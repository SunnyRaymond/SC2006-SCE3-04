from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "AI-Critique-Report.docx"
GREEN = "087F5B"
GREEN_DARK = "065F46"
INK = "0F172A"
SLATE = "475569"
SOFT = "F1F5F9"

PROMPT = (
    "Critique the attached use case diagram image with respect to how well it matches the attached "
    "PlugPlan SG Lab 1 Software Requirements Specification, version 1.0. Identify concrete mismatches, "
    "omissions, ambiguous actor responsibilities, incorrect UML semantics, and traceability gaps. For each "
    "finding, cite the relevant FR/NFR/Use Case identifier and explain its impact. Do not rewrite the diagram; "
    "return only the critique."
)

ORIGINAL_RESPONSE = r'''## Overall assessment

The diagram has strong top-level functional coverage: all eight use cases in the SRS catalogue appear inside the system boundary, and the primary human actors are associated with their expected journeys. The main weaknesses concern actor responsibility, overlapping use-case boundaries, and traceability rather than missing calculation details.

1. **UC-05 omits a supporting external actor that its specification may require.** UC-05 says the system repeats feasibility and ranking after the Driver changes plan inputs, and UC-05.EX.1 explicitly handles a route-refresh failure. A route refresh invokes OneMap under FR-06, while recalculation may also depend on charging data under FR-07. However, the diagram associates External Data Providers with UC-03 and UC-06 only. This hides an external dependency and makes the applicability of NFR-02 to UC-05 unclear. If UC-05 intentionally reuses the existing route and Official Snapshot, its route-refresh exception is inconsistent with that design; otherwise, the association is missing.

2. **The boundary between UC-04 and UC-05 is ambiguous for strategy changes.** UC-04.AC.1 allows the Driver to change the ranking strategy and see the order change, but FR-21 requires every input or strategy change to create a new immutable Plan Version, which is the stated responsibility of UC-05. The diagram presents UC-04 and UC-05 as independent use cases and explicitly asserts no `include` or `extend` relationship. This leaves it unclear whether changing strategy during comparison merely reorders the current results or must invoke the versioned recalculation behavior. The ambiguity could cause an implementation to bypass FR-21 and lose auditability.

3. **Combining OneMap and LTA DataMall into one actor weakens dependency traceability.** The SRS acknowledges this as a readability simplification, so it is not inherently invalid in an initial overview. Nevertheless, the two systems perform non-interchangeable roles: OneMap supplies place and route information for FR-06, whereas LTA DataMall supplies charging data for FR-07 and is subject to the specific freshness rule in NFR-03. A single actor association to UC-03 and UC-06 cannot show which provider participates in which step or which failure and freshness constraints apply. This limits the diagram’s usefulness for interface ownership and integration-test planning.

4. **UC-06’s name implies that fallback selection always occurs.** FR-22 requires manual refresh, whereas FR-23 requires a fallback only when the previously selected connector becomes infeasible. UC-06.AC.1 also confirms that a refreshed plan may require no fallback. The combined label “Refresh Plan and Select Fallback” therefore overstates the unconditional behavior of the use case. This is a semantic and acceptance-test ambiguity: a successful refresh could be incorrectly treated as incomplete unless a fallback is selected.

5. **UC-01 obscures different account privileges for Driver and Moderator.** Both actors connect to the single “Access Account” use case. In the SRS, FR-01 permits public creation of a Driver account, FR-02 authenticates either role, and UC-01 states that Moderator accounts are provisioned rather than publicly registered. The diagram does not distinguish these responsibilities and can therefore be read as permitting a Moderator to participate in account creation. The prose resolves the ambiguity, but the diagram alone does not communicate the authorization boundary.

6. **UC-08 leaves the source of “current official connector status” unresolved.** UC-08 step 2 requires the Moderator to see the current official connector status, but the diagram does not associate External Data Providers with UC-08. If UC-08 retrieves current LTA DataMall information, the supporting actor is omitted and NFR-02/NFR-03 should apply. If it uses only an already stored Official Snapshot, “current” is too strong and the snapshot age should be specified. This affects the evidence provenance on which FR-25 moderation decisions are based.

7. **The NFR trace is too coarse to resolve the preceding ambiguities.** The SRS maps every FR range to one or more use cases, so functional coverage is complete. In contrast, NFR-01 through NFR-10 are collectively mapped to “All applicable.” Important mappings therefore remain implicit, including NFR-02/NFR-03 for UC-03, UC-05, and UC-06; NFR-06 for UC-01 and UC-08; and NFR-07 for the profile-deletion path in UC-02. A Use Case Diagram does not normally contain NFRs, so their absence from the picture is not a UML defect, but the accompanying traceability does not compensate with sufficiently specific links.

8. **Association routing reduces readability without changing the underlying semantics.** The Driver fan-out and the long Moderator-to-UC-01 association cross several other lines near UC-05 and UC-07. The endpoints can be traced at full size, but they are easy to misread in a scaled report or presentation. The unexplained blue, green, and yellow use-case fills add another mild ambiguity because no legend states whether color represents actor, priority, or function. This is principally a presentation problem, not evidence of an incorrect actor association.

The absence of separate top-level use cases for arrival-SOC calculation, connector filtering, target-SOC calculation, charging-time estimation, cost estimation, and ranking under FR-08 through FR-19 is **not** an omission: these are internal system behaviours within UC-03 rather than independent actor goals. Likewise, omitting an `include` relationship from every authenticated use case to UC-01 is acceptable because authentication is expressed as a precondition. The general decision to avoid `include` and `extend` relationships is therefore defensible for an initial model, although it does not resolve the specific UC-04/UC-05 and UC-06 boundary ambiguities above.'''

SELECTED_QUOTE = (
    "The boundary between UC-04 and UC-05 is ambiguous for strategy changes. UC-04.AC.1 allows the Driver "
    "to change the ranking strategy and see the order change, but FR-21 requires every input or strategy "
    "change to create a new immutable Plan Version, which is the stated responsibility of UC-05. The diagram "
    "presents UC-04 and UC-05 as independent use cases and explicitly asserts no `include` or `extend` "
    "relationship. This leaves it unclear whether changing strategy during comparison merely reorders the "
    "current results or must invoke the versioned recalculation behavior. The ambiguity could cause an "
    "implementation to bypass FR-21 and lose auditability."
)

JUSTIFICATION = (
    "This critique was selected because it identifies a concrete consistency risk between UC-04.AC.1 and "
    "FR-21 rather than a cosmetic omission. If the boundary stays ambiguous, different team members could "
    "implement a strategy change either as an in-place reorder or as an immutable recalculation, producing "
    "incompatible tests and weakening auditability. The finding gives the team a precise decision to make in "
    "a later version: assign strategy changes to one use-case boundary and update the requirement, flow, and "
    "diagram together."
)


def shade(paragraph, fill: str) -> None:
    props = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_info_table(doc: Document) -> None:
    rows = [
        ("Version", "1.0"),
        ("Date", "4 September 2026"),
        ("Prepared for", "SC2006 Team SCE3-04"),
        ("Independent reviewer", "OpenAI Codex subagent"),
        ("Reviewed SRS", "PlugPlan-SG-Lab1-SRS.pdf, version 1.0"),
        ("Reviewed diagram", "Use-Case-Diagram.drawio.png"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.65)
    table.columns[1].width = Inches(4.9)
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        shade(table.cell(i, 0).paragraphs[0], "E2E8F0")
        table.cell(i, 0).paragraphs[0].runs[0].bold = True
        table.cell(i, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(GREEN_DARK)


def add_raw_response(doc: Document) -> None:
    for line in ORIGINAL_RESPONSE.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3 if line else 1)
        p.paragraph_format.keep_together = True
        if line:
            shade(p, "F8FAFC")
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(8.3)
            run.font.color.rgb = RGBColor.from_string(INK)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.68)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].font.color.rgb = RGBColor.from_string(INK)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for name, size, color in [("Title", 28, INK), ("Heading 1", 18, GREEN_DARK), ("Heading 2", 13, GREEN_DARK)]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True
        styles[name].font.color.rgb = RGBColor.from_string(color)

    quote_style = styles.add_style("Selected Quote", WD_STYLE_TYPE.PARAGRAPH)
    quote_style.font.name = "Arial"
    quote_style.font.size = Pt(10.5)
    quote_style.font.italic = True
    quote_style.font.color.rgb = RGBColor.from_string(SLATE)
    quote_style.paragraph_format.left_indent = Inches(0.3)
    quote_style.paragraph_format.right_indent = Inches(0.3)
    quote_style.paragraph_format.space_before = Pt(8)
    quote_style.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = "SC2006 | Team SCE3-04 | PlugPlan SG"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(SLATE)
    add_page_number(section.footer.paragraphs[0])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(78)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PLUGPLAN SG")
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GREEN)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("AI Critique Report")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Lab 1 - Requirements Elicitation")
    sr.font.name = "Arial"
    sr.font.size = Pt(18)
    sr.font.color.rgb = RGBColor.from_string(SLATE)
    doc.add_paragraph()
    add_info_table(doc)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(36)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run("The prompt and original AI response below are preserved as the critique evidence for the frozen version 1.0 inputs.")
    nr.font.name = "Arial"
    nr.font.size = Pt(10)
    nr.font.italic = True
    nr.font.color.rgb = RGBColor.from_string(SLATE)

    doc.add_page_break()
    doc.add_heading("1. Review Inputs and Baseline Freeze", level=1)
    doc.add_paragraph(
        "The reviewer received the final SRS version 1.0 and the final PNG export of the editable draw.io Use Case Diagram. "
        "No requirement or diagram changes were made after the critique request, preserving the Lab 1 evidence chain."
    )

    doc.add_heading("2. Prompt Sent to the AI Reviewer", level=1)
    prompt_p = doc.add_paragraph(PROMPT)
    prompt_p.paragraph_format.left_indent = Inches(0.22)
    prompt_p.paragraph_format.right_indent = Inches(0.22)
    shade(prompt_p, "E0F2FE")

    doc.add_heading("3. Original AI Response (Verbatim)", level=1)
    intro = doc.add_paragraph("The following response is reproduced verbatim, including its original Markdown markers.")
    intro.runs[0].italic = True
    add_raw_response(doc)

    doc.add_page_break()
    doc.add_heading("4. Selected Critique", level=1)
    doc.add_paragraph(SELECTED_QUOTE, style="Selected Quote")

    doc.add_heading("5. Team Justification for the Selection", level=1)
    doc.add_paragraph(JUSTIFICATION)

    doc.add_heading("6. Disposition", level=1)
    doc.add_paragraph(
        "This report records review feedback; it does not silently revise the artifacts that were reviewed. The selected "
        "finding and the remaining critique items should be assessed by the team before the next requirements version. "
        "Any accepted change must update the relevant FR, Use Case description, traceability entry, and UML diagram together."
    )

    doc.add_heading("7. AI Assistance Disclosure", level=1)
    doc.add_paragraph(
        "A separate OpenAI Codex subagent generated the original critique after receiving only the frozen SRS and Use Case "
        "Diagram. OpenAI Codex formatted this report and recorded the selected quotation and justification. The team remains "
        "responsible for checking the critique and deciding which findings to adopt."
    )

    doc.core_properties.title = "PlugPlan SG - Lab 1 AI Critique Report"
    doc.core_properties.subject = "AI critique of the frozen SRS and UML Use Case Diagram"
    doc.core_properties.author = "SC2006 Team SCE3-04"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
