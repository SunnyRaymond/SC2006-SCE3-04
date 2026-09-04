from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "PlugPlan-SG-Lab1-SRS.docx"

GREEN = "087F5B"
DARK_GREEN = "065F46"
INK = "0F172A"
SLATE = "475569"
LIGHT = "F1F5F9"
LINE = "CBD5E1"
WHITE = "FFFFFF"


functional_requirements = [
    ("FR-01", "High", "The system shall create a Driver account when a user submits a unique email address and a password that satisfies the configured password policy."),
    ("FR-02", "High", "The system shall authenticate a registered user and establish a session containing the user's Driver or Moderator role."),
    ("FR-03", "High", "The system shall allow a Driver to create, view, update, and delete one EV profile containing usable battery capacity, energy consumption, supported plug types, AC/DC charging limits, and reserve SOC."),
    ("FR-04", "High", "The system shall reject an EV profile containing a missing or out-of-range required value and display a field-specific validation message."),
    ("FR-05", "High", "The system shall create a charging-plan request from a selected EV profile, origin, destination, current SOC, departure time, maximum charging duration, maximum detour, ranking strategy, and optional manual target SOC."),
    ("FR-06", "High", "The system shall use OneMap to resolve submitted places and obtain route distance and duration, or report that a place is ambiguous or a route is unavailable."),
    ("FR-07", "High", "The system shall retrieve the latest usable charging-station and connector snapshot through the LTA DataMall backend adapter."),
    ("FR-08", "High", "The system shall estimate the arrival SOC at each candidate station from route distance, current SOC, usable battery capacity, and declared vehicle consumption."),
    ("FR-09", "High", "The system shall exclude a candidate station when the estimated battery energy on arrival is below zero."),
    ("FR-10", "High", "The system shall exclude a candidate station when it has no connector matching the vehicle's supported plug types."),
    ("FR-11", "High", "The system shall exclude a candidate station when it is closed during the planned charging window."),
    ("FR-12", "High", "The system shall exclude a candidate station when its route detour exceeds the Driver's maximum detour."),
    ("FR-13", "High", "The system shall calculate the automatic target SOC required to travel from the charger to the destination while retaining the configured reserve SOC."),
    ("FR-14", "High", "The system shall allow the Driver to replace the automatic target with a valid manual target while continuing to display the automatic target for comparison."),
    ("FR-15", "High", "The system shall estimate required charging energy and charging time using the lower of connector power and the applicable vehicle AC/DC charging limit."),
    ("FR-16", "High", "The system shall classify a candidate as dwell-infeasible when its estimated charging time exceeds the submitted maximum charging duration."),
    ("FR-17", "High", "The system shall estimate charging cost for supported per-kWh or per-hour prices and display 'Price unavailable' for a missing or unsupported price."),
    ("FR-18", "High", "The system shall rank feasible candidates using the selected Fastest overall journey, Cheapest charging, Availability-first, Minimum detour, or Balanced preset."),
    ("FR-19", "High", "The system shall display two or three shortlisted recommendations with arrival SOC, target SOC, required energy, charging time, estimated cost, ranking explanation, source timestamps, and data-state label."),
    ("FR-20", "Medium", "The system shall allow the Driver to compare two or three shortlisted recommendations side by side."),
    ("FR-21", "High", "The system shall recalculate a plan after the Driver changes an input or strategy and store the result as a new immutable Plan Version."),
    ("FR-22", "High", "The system shall allow the Driver to manually refresh a saved plan using the latest usable external-data snapshot."),
    ("FR-23", "High", "The system shall identify the next-ranked feasible alternative when a refreshed plan makes the previously selected connector infeasible, or state that no feasible fallback exists."),
    ("FR-24", "Medium", "The system shall allow a Driver to submit a time-limited charger issue report containing a station or connector, category, note, and optional image."),
    ("FR-25", "Medium", "The system shall allow only a Moderator to verify, reject, merge, resolve, or expire an issue report while recording the decision reason and timestamp."),
]

nonfunctional_requirements = [
    ("NFR-01", "Performance", "With cached provider data, 95% of 100 recommendation requests under 20 concurrent users shall complete within 3.0 seconds on the team's documented reference environment.", "Load test results and percentile report."),
    ("NFR-02", "Failure response", "When an external provider times out or returns an error, the system shall display a degraded-state result or explicit failure within 5 seconds.", "Provider-failure integration tests."),
    ("NFR-03", "Freshness", "The system shall label LTA data older than 10 minutes as Stale and shall display observedAt and retrievedAt wherever a recommendation is shown.", "Timestamp boundary tests at 9:59 and 10:00 minutes."),
    ("NFR-04", "Calculation correctness", "Approved fixtures shall produce 100% correct feasibility decisions and numerical results within +/-0.1 SOC percentage point, +/-0.1 kWh, +/-1 minute, and +/-S$0.01.", "Fixture-based unit and acceptance tests."),
    ("NFR-05", "Security", "All authenticated traffic shall use TLS 1.2 or later; passwords shall be stored only as salted one-way hashes; no provider secret shall appear in client code, repository history, or application logs.", "Configuration review, secret scan, and transport inspection."),
    ("NFR-06", "Authorization", "Every Moderator endpoint shall reject an authenticated Driver with HTTP 403, as verified by an automated test for each protected operation.", "Automated role-access tests."),
    ("NFR-07", "Privacy", "After a Driver deletes an EV profile or saved plan, it shall become inaccessible through the application within 60 seconds; unsaved origin and destination inputs shall not be retained.", "Deletion and storage inspection tests."),
    ("NFR-08", "Accessibility", "The nine core mockup frames shall have zero serious or critical automated accessibility violations, and the implemented core journeys shall support keyboard-only completion.", "Automated audit plus keyboard walkthrough."),
    ("NFR-09", "Usability", "At least 4 of 5 first-time representative users shall create a plan and compare two recommendations without assistance within 3 minutes.", "Moderated usability test with timing and success records."),
    ("NFR-10", "Reliability and testability", "The deterministic fixture-based recommendation journey shall complete successfully in 100 consecutive runs, and energy, feasibility, cost, and ranking modules shall achieve at least 80% branch coverage.", "Repeat-run test and coverage report."),
]

data_dictionary = [
    ("User", "A person who accesses PlugPlan SG.", "userId; unique email; passwordHash; role {DRIVER, MODERATOR}; status", "A Driver owns an EV Profile, Charging Plans, and Issue Reports. A Moderator records Moderation Decisions."),
    ("EV Profile", "Planning assumptions for one Driver's vehicle.", "profileId; usableCapacityKWh 10-250; consumptionKWhPer100Km 5-50; plugTypes; maxACPowerKW 1-50; maxDCPowerKW 1-400; reserveSOCPercent 5-40", "Exactly one belongs to each Driver in the MVP and is referenced by Charging Plans."),
    ("Charging Plan", "A Driver's persistent charging-planning record.", "planId; driverId; profileId; status; createdAt", "Belongs to one Driver and contains one or more immutable Plan Versions."),
    ("Plan Version", "An immutable set of plan inputs and results.", "versionId; planId; origin; destination; currentSOCPercent 0-100; departureAt; maximumChargingMinutes 5-720; maximumDetourMinutes 0-120; strategy; optional manualTargetSOCPercent; createdAt", "Belongs to one Charging Plan; references one Official Snapshot and multiple Recommendations."),
    ("Charging Station", "A physical EV charging location.", "stationId; name; address; latitude; longitude; operator; operatingHours", "Contains one or more Connectors and may have Issue Reports."),
    ("Connector", "An individual charging interface at a station.", "connectorId; stationId; plugType; currentType {AC, DC}; maximumPowerKW; status; priceType; optional unitPrice", "Belongs to one Charging Station and is referenced by Recommendations and Issue Reports."),
    ("Official Snapshot", "A timestamped copy of external provider data used by a calculation.", "snapshotId; source; observedAt; retrievedAt; dataState {LIVE, CACHED, STALE, DEMO_FIXTURE}", "Supplies evidence to one or more Plan Versions."),
    ("Route Estimate", "A OneMap route result used for travel and energy calculations.", "routeId; origin; destination; distanceKm; durationMinutes; retrievedAt", "Used by a Plan Version to calculate arrival SOC, detour, and destination energy."),
    ("Recommendation", "A ranked, explained assessment of one connector for one Plan Version.", "recommendationId; versionId; connectorId; feasibility; arrivalSOC; automaticTargetSOC; requiredEnergyKWh; chargingMinutes; optional estimatedCost; rank; explanation", "Belongs to one Plan Version and refers to one Connector."),
    ("Ranking Strategy", "A fixed factor-weight preset.", "FASTEST; CHEAPEST; AVAILABILITY_FIRST; MINIMUM_DETOUR; BALANCED", "Selected by a Plan Version and applied to its Recommendations."),
    ("Issue Report", "A time-limited Driver report about a charging location or connector.", "reportId; reporterId; stationId; optional connectorId; category; note; optional imageReference; status; createdAt; expiresAt", "Submitted by one Driver and receives zero or more Moderation Decisions."),
    ("Report Status", "The lifecycle state of an Issue Report.", "PENDING; VERIFIED; REJECTED; MERGED; RESOLVED; EXPIRED", "Constrains valid transitions for one Issue Report."),
    ("Moderation Decision", "An auditable action taken by a Moderator.", "decisionId; reportId; moderatorId; action; reason; createdAt", "Belongs to one Issue Report and one Moderator."),
]

use_cases = [
    {
        "id": "UC-01", "name": "Access Account", "actor": "Driver or Moderator", "priority": "High", "frequency": "At the start of an authenticated session",
        "description": "A new Driver registers, or an existing Driver or Moderator signs in and reaches the role-appropriate workspace.",
        "pre": ["The user can access the PlugPlan SG web application.", "For sign-in, an active account exists."],
        "post": ["An authenticated session contains the user's role.", "The user is directed to the appropriate Driver or Moderator workspace."],
        "flow": ["The user opens the account access screen.", "The user enters an email address and password.", "The system validates the credentials.", "The system creates an authenticated session containing the user's role.", "The system displays the role-appropriate workspace."],
        "alts": ["UC-01.AC.1 - A new Driver selects Create a Driver account, enters a unique email and compliant password, and the system creates the account before continuing at step 4.", "UC-01.AC.2 - A signed-in user ends the session; the system invalidates the session and returns to the access screen."],
        "exceptions": ["UC-01.EX.1 - If credentials are invalid, the system displays a generic error and does not reveal whether the email exists.", "UC-01.EX.2 - If the account is inactive, the system denies access and provides a support instruction."],
        "special": "NFR-05 and NFR-06 apply.", "assumptions": "Moderator accounts are provisioned by the team; public registration creates Driver accounts only.", "trace": "FR-01, FR-02; D-00"
    },
    {
        "id": "UC-02", "name": "Manage EV Profile", "actor": "Driver", "priority": "High", "frequency": "Initial setup and occasional updates",
        "description": "The Driver records validated vehicle and planning assumptions used by subsequent charging plans.",
        "pre": ["The Driver is authenticated."],
        "post": ["The valid EV Profile is stored.", "Existing Plan Versions retain their original assumptions."],
        "flow": ["The Driver opens Vehicle Profile and Preferences.", "The system displays the current values or an empty form.", "The Driver enters vehicle capacity, consumption, plug types, charging limits, reserve, and planning preferences.", "The Driver selects Save Profile.", "The system validates the values and stores the profile.", "The system confirms that the profile was saved."],
        "alts": ["UC-02.AC.1 - The Driver edits an existing profile; future Plan Versions use the new values.", "UC-02.AC.2 - The Driver deletes the profile after confirming the destructive action."],
        "exceptions": ["UC-02.EX.1 - Missing or out-of-range values are rejected with field-specific messages.", "UC-02.EX.2 - If storage fails, the system keeps the entered values and displays a retry instruction."],
        "special": "Values must show units and must be described as planning assumptions.", "assumptions": "The MVP stores exactly one EV Profile per Driver.", "trace": "FR-03, FR-04; D-01"
    },
    {
        "id": "UC-03", "name": "Create Charging Plan", "actor": "Driver; supporting actor: External Data Providers", "priority": "High", "frequency": "Several times per week per active Driver",
        "description": "The Driver submits journey inputs and receives a saved Plan Version containing ranked, explained charging recommendations.",
        "pre": ["The Driver is authenticated.", "A valid EV Profile exists."],
        "post": ["A new Charging Plan and Plan Version are stored.", "Feasible recommendations or an explicit no-feasible-option result is displayed."],
        "flow": ["The Driver enters origin, destination, current SOC, departure time, maximum charging duration, maximum detour, and target choice.", "The system resolves the places and obtains route distance and duration from OneMap.", "The system obtains the latest usable station and connector snapshot through the LTA DataMall adapter.", "The system calculates reachability, compatibility, arrival SOC, target SOC, required energy, charging time, cost, and dwell fit.", "The system excludes infeasible candidates and ranks the remaining candidates using the selected fixed strategy.", "The system stores the Plan Version and displays two or three recommendations with explanations, timestamps, and data-state labels."],
        "alts": ["UC-03.AC.1 - The Driver selects a valid manual target; the system uses it while continuing to show the automatic target.", "UC-03.AC.2 - The Driver already has enough energy; the system offers continuation without a charging stop.", "UC-03.AC.3 - No candidate is feasible; the system explains the reasons and offers editable inputs."],
        "exceptions": ["UC-03.EX.1 - An ambiguous place requires Driver selection before route calculation.", "UC-03.EX.2 - If no usable route is available, the system preserves the inputs and displays an explicit failure.", "UC-03.EX.3 - If live charging data fails, the system may use a labeled cached or Demo Fixture snapshot."],
        "special": "NFR-01 to NFR-04 apply. A result is an estimate, not a reservation or guarantee.", "assumptions": "Reserve is 20%, charging efficiency is 90%, and targets above 80% show a taper warning.", "trace": "FR-05 to FR-19; D-02, D-03"
    },
    {
        "id": "UC-04", "name": "Review and Compare Recommendations", "actor": "Driver", "priority": "High", "frequency": "After each successful plan calculation",
        "description": "The Driver inspects recommendation evidence and compares two or three alternatives using a consistent metric order.",
        "pre": ["A Plan Version contains at least two feasible Recommendations."],
        "post": ["The Driver understands the trade-offs and may select one recommendation."],
        "flow": ["The system displays ranked recommendation cards and the selected strategy.", "The Driver reviews availability, detour, arrival SOC, target SOC, required energy, time, cost, freshness, issues, and explanation.", "The Driver selects two or three recommendations for comparison.", "The system displays the selected recommendations side by side using the same metric order.", "The Driver reviews scoring factors, missing data, and exclusions.", "The Driver selects a preferred recommendation or returns without selecting one."],
        "alts": ["UC-04.AC.1 - The Driver changes the fixed ranking strategy; the display explains why the order changes.", "UC-04.AC.2 - A price is missing; the system displays Price unavailable and does not label that option cheapest."],
        "exceptions": ["UC-04.EX.1 - If fewer than two feasible choices remain, comparison is unavailable and the system explains why."],
        "special": "Official data, PlugPlan estimates, and community reports must be visually distinguished.", "assumptions": "A maximum of three options can be compared at once.", "trace": "FR-18 to FR-20; D-03, D-04"
    },
    {
        "id": "UC-05", "name": "Reconfigure Charging Plan", "actor": "Driver", "priority": "High", "frequency": "As needed before selecting a charger",
        "description": "The Driver changes a plan input or strategy and receives a new immutable Plan Version.",
        "pre": ["An existing Charging Plan has at least one Plan Version."],
        "post": ["A new Plan Version is stored.", "The previous Plan Version remains unchanged and accessible."],
        "flow": ["The Driver opens Compare and Reconfigure for an existing plan.", "The system displays current timing, dwell, reserve, target, detour, and strategy values.", "The Driver changes one or more values.", "The Driver selects Recalculate.", "The system repeats feasibility and ranking using the new values.", "The system stores and displays the new Plan Version with an explanation of material ranking changes."],
        "alts": ["UC-05.AC.1 - The Driver resets the pending changes and keeps the current Plan Version.", "UC-05.AC.2 - A previously feasible option becomes infeasible and is labeled accordingly."],
        "exceptions": ["UC-05.EX.1 - If route refresh fails, the previous Plan Version remains visible and no partial version is stored.", "UC-05.EX.2 - Leaving with unsaved changes requires confirmation."],
        "special": "Plan Version immutability supports traceability.", "assumptions": "Reconfiguration is user-initiated; continuous background recalculation is outside scope.", "trace": "FR-05, FR-14, FR-18, FR-21; D-04"
    },
    {
        "id": "UC-06", "name": "Refresh Plan and Select Fallback", "actor": "Driver; supporting actor: External Data Providers", "priority": "High", "frequency": "Immediately before travel or after a visible data-age warning",
        "description": "The Driver manually requests newer data and reviews a fallback when the selected connector is no longer feasible.",
        "pre": ["A saved Plan Version exists.", "The Driver is authenticated."],
        "post": ["A refresh result is stored as a new Plan Version.", "The prior selection remains or a fallback result is clearly presented."],
        "flow": ["The Driver selects Refresh Data on a saved plan.", "The system requests the latest usable route and charging snapshot through its backend adapters.", "The system recalculates feasibility and ranking using the existing plan inputs.", "The system compares the refreshed result with the previous Plan Version.", "If the selected connector is unsuitable, the system explains the change and identifies the next-ranked feasible option.", "The Driver keeps the original selection when still feasible, selects the fallback, or edits the plan."],
        "alts": ["UC-06.AC.1 - The original selection remains feasible; the system confirms that no fallback is required.", "UC-06.AC.2 - No feasible fallback exists; the system provides reasons and editable inputs.", "UC-06.AC.3 - Only cached or stale data is available; the system labels the state and age."],
        "exceptions": ["UC-06.EX.1 - If no usable snapshot is available, the system displays a failure and keeps the previous Plan Version unchanged."],
        "special": "Refresh is manual. The system must not promise automatic alerts or guaranteed availability.", "assumptions": "The latest usable snapshot can be live, cached, stale, or a clearly labeled Demo Fixture.", "trace": "FR-07 to FR-19, FR-22, FR-23; D-03, D-04"
    },
    {
        "id": "UC-07", "name": "Submit Charger Issue", "actor": "Driver", "priority": "Medium", "frequency": "Occasionally after observing a site problem",
        "description": "The Driver submits a categorized, expiring community report associated with a station or connector.",
        "pre": ["The Driver is authenticated.", "A Charging Station or Connector is selected."],
        "post": ["A Pending Issue Report is stored with an expiry time.", "The Driver receives a report reference."],
        "flow": ["The Driver selects Report an Issue for a station or connector.", "The system pre-fills the station and connector context.", "The Driver selects a category, enters a note, and optionally adds an image.", "The system validates the input and checks for a similar active report.", "The Driver confirms submission.", "The system stores a Pending report, sets its expiry, and displays the reference."],
        "alts": ["UC-07.AC.1 - A similar report exists; the Driver views it instead of submitting a duplicate.", "UC-07.AC.2 - The Driver opens the operator's official reporting channel for repair or enforcement."],
        "exceptions": ["UC-07.EX.1 - An invalid or oversized image is rejected without discarding other inputs.", "UC-07.EX.2 - If the submission rate limit is exceeded, the system displays when another report may be submitted."],
        "special": "Images must include a privacy reminder. Community reports must not look like official status.", "assumptions": "Default report expiry is 24 hours unless moderated earlier.", "trace": "FR-24; D-05"
    },
    {
        "id": "UC-08", "name": "Moderate Charger Reports", "actor": "Moderator", "priority": "Medium", "frequency": "Daily during the prototype demonstration period",
        "description": "A Moderator reviews community evidence and records a justified report-lifecycle decision.",
        "pre": ["The Moderator is authenticated.", "At least one Issue Report exists."],
        "post": ["A valid Moderation Decision is appended to the audit history.", "The Issue Report status reflects the decision."],
        "flow": ["The Moderator opens the restricted report queue and selects a report.", "The system displays the report, evidence, related reports, current official connector status, and audit history.", "The Moderator reviews the provenance and enters a decision reason.", "The Moderator selects Verify, Reject, Merge, Resolve, or Expire.", "The system requests confirmation for destructive or consolidating actions.", "The system records the actor, action, reason, timestamp, and resulting status."],
        "alts": ["UC-08.AC.1 - Evidence is inconclusive; the Moderator keeps the report Pending.", "UC-08.AC.2 - An official status supersedes the report; the Moderator resolves it with that reason.", "UC-08.AC.3 - The Moderator selects a valid target report and merges a duplicate."],
        "exceptions": ["UC-08.EX.1 - If another Moderator changed the report, the system rejects the stale update and requests a refresh.", "UC-08.EX.2 - A missing decision reason prevents submission."],
        "special": "NFR-06 applies to every moderation operation.", "assumptions": "No appeal workflow is included in the MVP.", "trace": "FR-02, FR-25; D-06"
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header=True, font_size=8.5):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(font_size)
                    r.font.color.rgb = RGBColor.from_string(INK)
            if header and ri == 0:
                set_cell_shading(cell, DARK_GREEN)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor.from_string(WHITE)
            elif ri % 2 == 0:
                set_cell_shading(cell, "F8FAFC")


def add_table(doc, headers, rows, widths=None, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=WHITE, size=font_size)
        set_cell_shading(table.rows[0].cells[i], DARK_GREEN)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    style_table(table, header=True, font_size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc, items):
    for index, item in enumerate(items, start=1):
        # Use explicit numbering so each Use Case flow reliably restarts at 1
        # across Word, LibreOffice, and PDF exports.
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(f"{index}. ")
        p.add_run(item)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_section(section, landscape=False):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Inches(0.72)
        section.right_margin = Inches(0.72)
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.clear()
    header.text = "SC2006 | Team SCE3-04 | PlugPlan SG"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in header.runs:
        r.font.name = "Arial"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(SLATE)
    footer = section.footer.paragraphs[0]
    footer.clear()
    add_page_number(footer)


def add_section(doc, landscape=False):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape)
    return section


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PLUGPLAN SG")
    r.font.name = "Arial"
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GREEN)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("Software Requirements Specification")
    r.font.name = "Arial"
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Lab 1 - Requirements Elicitation")
    r.font.name = "Arial"
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(SLATE)
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    rows = [
        ("Version", "1.0 - requirements frozen for AI critique"),
        ("Date", "4 September 2026"),
        ("Prepared for", "SC2006 Team SCE3-04"),
        ("Product", "PlugPlan SG"),
        ("Document status", "Initial Lab 1 baseline"),
    ]
    for i, (a, b) in enumerate(rows):
        set_cell_text(table.rows[i].cells[0], a, bold=True, color=DARK_GREEN, size=10)
        set_cell_text(table.rows[i].cells[1], b, size=10)
        table.rows[i].cells[0].width = Inches(1.6)
        table.rows[i].cells[1].width = Inches(4.5)
        if i % 2 == 0:
            set_cell_shading(table.rows[i].cells[0], LIGHT)
            set_cell_shading(table.rows[i].cells[1], LIGHT)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(56)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Estimates only. PlugPlan SG does not reserve chargers, process payment, control charging hardware, or guarantee availability.")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string(SLATE)
    doc.add_page_break()


def add_use_case(doc, uc):
    doc.add_heading(f"{uc['id']} {uc['name']}", level=2)
    meta = [
        ("Primary actor(s)", uc["actor"]),
        ("Description", uc["description"]),
        ("Priority", uc["priority"]),
        ("Frequency", uc["frequency"]),
        ("Traceability", uc["trace"]),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in meta:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, color=DARK_GREEN, size=8.5)
        set_cell_text(cells[1], value, size=8.5)
        set_cell_shading(cells[0], LIGHT)
    doc.add_heading("Preconditions", level=3)
    add_bullets(doc, uc["pre"])
    doc.add_heading("Postconditions", level=3)
    add_bullets(doc, uc["post"])
    doc.add_heading("Normal Flow of Events", level=3)
    add_numbered(doc, uc["flow"])
    doc.add_heading("Alternative Flows", level=3)
    add_bullets(doc, uc["alts"])
    doc.add_heading("Exceptions", level=3)
    add_bullets(doc, uc["exceptions"])
    table2 = doc.add_table(rows=3, cols=2)
    table2.style = "Table Grid"
    for i, (label, value) in enumerate([
        ("Included Use Cases", "None in this initial model."),
        ("Special Requirements", uc["special"]),
        ("Assumptions and Notes", uc["assumptions"]),
    ]):
        set_cell_text(table2.rows[i].cells[0], label, bold=True, color=DARK_GREEN, size=8.5)
        set_cell_text(table2.rows[i].cells[1], value, size=8.5)
        set_cell_shading(table2.rows[i].cells[0], LIGHT)
    if uc != use_cases[-1]:
        doc.add_page_break()


def build():
    doc = Document()
    configure_section(doc.sections[0], landscape=False)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string(INK)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    for level, size in [(1, 20), (2, 15), (3, 11)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(DARK_GREEN if level < 3 else INK)
        style.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True
    add_title_page(doc)

    doc.add_heading("Document Control", level=1)
    add_table(doc, ["Version", "Date", "Change", "Owner"], [["1.0", "4 September 2026", "Initial Lab 1 requirements baseline used for the AI critique.", "Team SCE3-04"]], widths=[0.7, 1.2, 3.8, 1.2], font_size=9)
    doc.add_heading("Contribution and AI Disclosure", level=2)
    doc.add_paragraph("This package was assembled in the repository under Raymond Guo's direction with drafting, diagram, mockup, formatting, and independent critique assistance from OpenAI Codex. Every team member must review, understand, and approve the technical content before submission. The complete team-member contribution allocation must be recorded in the accompanying Contributions.md file.")
    doc.add_heading("Document Map", level=2)
    add_table(doc, ["Section", "Contents"], [
        ("1-3", "Purpose, scope, users, product context, and external interfaces"),
        ("4", "Functional Requirements"),
        ("5", "Non-Functional Requirements"),
        ("6", "Data Dictionary"),
        ("7", "Initial Use Case Diagram and Use Case Descriptions"),
        ("8", "UI Mockups and HCI requirements"),
        ("9", "Traceability and requirements decisions"),
    ], widths=[1, 5.8], font_size=9)

    doc.add_heading("1. Introduction", level=1)
    doc.add_heading("1.1 Purpose", level=2)
    doc.add_paragraph("This Software Requirements Specification defines the Lab 1 baseline for PlugPlan SG. It records atomic and verifiable requirements, important data terms, an initial UML Use Case Model, UI Mockups, and traceability links for later design, implementation, and testing.")
    doc.add_heading("1.2 Product Scope", level=2)
    doc.add_paragraph("PlugPlan SG is a responsive web-based decision-support application for Electric Vehicle (EV) drivers planning a charging stop in Singapore. A Driver provides an EV Profile, origin, destination, current State of Charge (SOC), departure time, maximum charging duration, and planning preferences. The system combines route estimates from OneMap with charging-station and connector data obtained through an LTA DataMall backend adapter. It identifies feasible connectors, estimates arrival SOC, recommends the minimum useful target SOC with a visible reserve, estimates charging time and cost, and ranks alternatives using fixed and explainable strategies. Drivers can compare alternatives, reconfigure a plan, manually refresh external data, select a fallback, and submit charger issue reports. Moderators review and resolve those reports.")
    doc.add_heading("1.3 Intended Audience", level=2)
    add_bullets(doc, ["Team SCE3-04 members implementing and testing the application.", "The SC2006 Lab Supervisor and Teaching Assistant reviewing Lab 1.", "Future maintainers tracing requirements to analysis, design, code, and tests."])
    doc.add_heading("1.4 Document Conventions", level=2)
    add_bullets(doc, ["Functional Requirements use unique FR identifiers and the normative word shall.", "Non-Functional Requirements use NFR identifiers and include a verification method.", "High priority means necessary for the main decision journey; Medium priority is required for the issue-reporting support journey.", "Terms defined in the Data Dictionary are capitalized when referring to the domain concept."])
    doc.add_heading("1.5 References", level=2)
    add_bullets(doc, ["SC2006 Lab 1 Manual, local course copy.", "SC2006 Frequently Asked Questions v4.1, local course copy.", "SC2006 supplied SRS Template and Use Case Template.", "PlugPlan SG Comprehensive Internal Proposal, scope update dated 4 September 2026.", "OneMap Search and Routing API documentation.", "LTA DataMall EV charging data documentation."])

    doc.add_heading("2. Overall Description", level=1)
    doc.add_heading("2.1 Product Perspective", level=2)
    doc.add_paragraph("PlugPlan SG is a new, self-contained web application. The browser communicates only with the PlugPlan backend. Backend adapters normalize data from OneMap and LTA DataMall before the recommendation engine uses it. A persistent store keeps accounts, the single MVP EV Profile per Driver, Charging Plans and immutable Plan Versions, official-data snapshots, Issue Reports, and Moderation Decisions.")
    doc.add_heading("2.2 User Classes and External Actors", level=2)
    add_table(doc, ["Actor", "Type", "Goal or responsibility"], [
        ("Driver", "Primary human actor", "Maintains an EV Profile, creates and compares plans, refreshes a plan, selects a fallback, and reports an issue."),
        ("Moderator", "Primary human actor", "Reviews evidence and verifies, rejects, merges, resolves, or expires charger Issue Reports."),
        ("External Data Providers", "Supporting system role", "OneMap resolves places and routes; LTA DataMall supplies station, connector, status, speed, price, and timestamp data."),
    ], widths=[1.25, 1.35, 4.1], font_size=9)
    doc.add_heading("2.3 In Scope", level=2)
    add_bullets(doc, ["Responsive Driver and Moderator web interfaces.", "One EV Profile per Driver in the MVP.", "Singapore journey planning using origin, destination, current SOC, departure time, dwell time, and maximum detour.", "Arrival-SOC, minimum useful target, charging-energy, charging-time, cost, and dwell-fit estimates.", "Five fixed strategies: Fastest overall journey, Cheapest charging, Availability-first, Minimum detour, and Balanced.", "Manual plan refresh, fallback explanation, issue reporting, and moderation.", "Live, Cached, Stale, and Demo Fixture data states."])
    doc.add_heading("2.4 Explicitly Out of Scope", level=2)
    add_bullets(doc, ["Official reservation, bay enforcement, payment, wallet, subscriptions, remote charger activation, or charger control.", "Continuous background monitoring, automatic alerts, or guaranteed future availability.", "Community queues, waitlists, offers, check-ins, or charging-session handovers.", "Traffic or carpark integration, operator account integration, native mobile applications, and in-car systems.", "Machine Learning availability prediction or a numeric reliability probability."])
    doc.add_heading("2.5 Fixed Requirements Decisions", level=2)
    add_table(doc, ["Decision", "Lab 1 baseline"], [
        ("Default reserve SOC", "20%; configurable from 5% to 40%."),
        ("Charging efficiency", "90%; always disclosed as an assumption."),
        ("High-SOC warning", "Displayed when a target exceeds 80%."),
        ("Fastest strategy", "Earliest estimated arrival at the final destination after route and charging time."),
        ("Unknown price", "Eligible for non-price strategies but never labelled the cheapest."),
        ("Issue expiry", "24 hours unless a Moderator acts earlier."),
        ("Data staleness", "LTA data older than 10 minutes is labelled Stale."),
    ], widths=[2, 4.7], font_size=9)

    doc.add_heading("3. External Interface Requirements", level=1)
    doc.add_heading("3.1 User Interface", level=2)
    add_bullets(doc, ["Every form uses persistent labels and visible units.", "Every calculated result is labelled as an estimate.", "Official data, PlugPlan calculations, and community reports use distinct text labels and visual treatments.", "A map is paired with a readable list alternative.", "Errors are placed next to the affected field and preserve user-entered values.", "The primary Driver journey has desktop and mobile mockups; moderation is desktop-first."])
    doc.add_heading("3.2 Software Interfaces", level=2)
    add_table(doc, ["Interface", "Input to PlugPlan", "Failure behavior"], [
        ("OneMap Search and Routing", "Resolved coordinates, driving distance, and route duration.", "Preserve inputs and report ambiguity, timeout, or unavailable route."),
        ("LTA DataMall adapter", "Station, connector, status, plug type, power, price, observed time, and retrieval time where available.", "Use a clearly labelled cached, stale, or Demo Fixture snapshot, or report that no usable data exists."),
        ("PlugPlan persistent store", "Accounts, profiles, plans, versions, reports, decisions, and snapshots.", "Do not store a partial Plan Version or Moderation Decision."),
    ], widths=[1.8, 3, 2.1], font_size=8.5)
    doc.add_heading("3.3 Communications and Security", level=2)
    doc.add_paragraph("Browser-to-backend and backend-to-provider communication uses HTTPS. Provider credentials remain server-side. The UI must never call a provider directly or expose provider secrets.")

    doc.add_heading("4. Functional Requirements", level=1)
    doc.add_paragraph("Each requirement is atomic enough to be verified independently. Inputs and outputs are made explicit in the requirement wording or linked Use Case.")
    add_table(doc, ["ID", "Priority", "Requirement"], functional_requirements, widths=[0.7, 0.8, 5.3], font_size=8.2)

    doc.add_heading("5. Non-Functional Requirements", level=1)
    add_table(doc, ["ID", "Quality", "Requirement", "Verification"], nonfunctional_requirements, widths=[0.65, 1.2, 3.45, 1.5], font_size=7.8)

    doc.add_heading("6. Data Dictionary", level=1)
    doc.add_paragraph("The dictionary defines important domain terms, principal attributes and constraints, and relationships. Detailed database types and keys will be refined during analysis and design.")
    add_table(doc, ["Term", "Definition", "Important attributes and constraints", "Relationships"], data_dictionary, widths=[1.1, 1.6, 2.7, 1.4], font_size=7.4)

    add_section(doc, landscape=True)
    doc.add_heading("7. Initial Use Case Model", level=1)
    doc.add_heading("7.1 Initial UML Use Case Diagram", level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Keep the figure and its caption together on one landscape page.
    p.add_run().add_picture(str(ROOT / "Use-Case-Diagram.drawio.png"), width=Inches(8.2))
    cap = doc.add_paragraph("Figure 1. PlugPlan SG initial Use Case Diagram. The two provider systems are shown as one supporting actor role to keep this initial model readable.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(8.5)

    add_section(doc, landscape=False)
    doc.add_heading("7.2 Use Case Catalogue", level=2)
    add_table(doc, ["ID", "Use Case", "Primary actor", "Success outcome", "FR trace"], [
        ("UC-01", "Access Account", "Driver / Moderator", "An authenticated role-aware session is created.", "FR-01 to FR-02"),
        ("UC-02", "Manage EV Profile", "Driver", "A valid single MVP EV Profile is stored.", "FR-03 to FR-04"),
        ("UC-03", "Create Charging Plan", "Driver", "A Plan Version and recommendations or an explicit no-feasible result are stored.", "FR-05 to FR-19"),
        ("UC-04", "Review and Compare Recommendations", "Driver", "Two or three alternatives are explained and compared.", "FR-18 to FR-20"),
        ("UC-05", "Reconfigure Charging Plan", "Driver", "A recalculated immutable Plan Version is stored.", "FR-21"),
        ("UC-06", "Refresh Plan and Select Fallback", "Driver", "Current data is applied and a fallback or no-fallback result is shown.", "FR-22 to FR-23"),
        ("UC-07", "Submit Charger Issue", "Driver", "A time-limited Pending Issue Report is stored.", "FR-24"),
        ("UC-08", "Moderate Charger Reports", "Moderator", "A justified lifecycle decision is recorded.", "FR-25"),
    ], widths=[0.6, 1.6, 1.1, 2.8, 0.9], font_size=8)
    doc.add_page_break()
    doc.add_heading("7.3 Use Case Descriptions", level=2)
    doc.add_paragraph("The normal flow for every Use Case contains no more than six steps, following the Lab 1 rule of thumb.")
    for uc in use_cases:
        add_use_case(doc, uc)

    add_section(doc, landscape=True)
    doc.add_heading("8. UI Mockups", level=1)
    doc.add_paragraph("The editable HTML and standalone UI-Mockups.pdf contain the same nine frames. The mockups intentionally expose normal, error, empty, stale-data, and fallback states because these states refine requirements rather than merely illustrate styling.")
    images = [
        ("D-00 Account Access", "D-00-Account-Access.png", "UC-01; role-aware account access"),
        ("D-01 Vehicle Profile and Preferences", "D-01-Vehicle-Profile.png", "UC-02; field validation and profile dependency warning"),
        ("D-02 Plan a Charging Stop", "D-02-Plan-Charging-Stop.png", "UC-03; journey inputs, automatic target, route error"),
        ("D-03 Charging Recommendations", "D-03-Recommendations.png", "UC-03, UC-04, UC-06; ranking, provenance, no-feasible and fallback states"),
        ("D-04 Compare and Reconfigure", "D-04-Compare-Reconfigure.png", "UC-04, UC-05; what-if inputs, immutable versioning, missing price"),
        ("D-05 Report a Charger Issue", "D-05-Report-Issue.png", "UC-07; duplicate, validation, privacy and expiry states"),
        ("D-06 Moderator Report Review", "D-06-Moderator-Review.png", "UC-08; restricted actions, audit trail and concurrent update"),
        ("M-01 Plan a Charging Stop - Mobile", "M-01-Plan-Mobile.png", "Responsive companion for UC-03"),
        ("M-02 Charging Recommendations - Mobile", "M-02-Recommendations-Mobile.png", "Responsive companion for UC-03, UC-04 and UC-06"),
    ]
    for idx, (title, filename, coverage) in enumerate(images, 1):
        if idx > 1:
            doc.add_page_break()
        doc.add_heading(title, level=2)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        is_mobile = filename.startswith("M-")
        p.add_run().add_picture(str(ROOT / "ui-mockups" / filename), height=Inches(5.85 if is_mobile else 5.75))
        cap = doc.add_paragraph(f"Figure {idx + 1}. {coverage}.")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(8.5)
    doc.add_page_break()
    doc.add_heading("8.1 HCI and Accessibility Requirements Elicited from the Mockups", level=2)
    add_bullets(doc, ["Use explicit labels rather than placeholder-only fields.", "Show units for SOC, kWh, kW, km, minutes, and S$.", "Use SGT for time labels.", "Provide visible keyboard focus and a logical tab order in the implementation.", "Use icon and text, not colour alone, for status.", "Preserve entered values after provider errors.", "Explain disabled or unavailable actions.", "Require confirmation for profile deletion and destructive moderation actions.", "Keep metric order consistent between recommendation and comparison views.", "Keep data freshness, provenance, and estimate disclaimers visible on mobile.", "Pair every map with a list or text alternative.", "Use a 44 by 44 pixel target for touch controls where practical."])

    add_section(doc, landscape=False)
    doc.add_heading("9. Traceability and Requirements Decisions", level=1)
    doc.add_heading("9.1 Requirements-to-Use-Case-to-Screen Trace", level=2)
    add_table(doc, ["Requirement range", "Use Case(s)", "Mockup(s)", "Primary verification focus"], [
        ("FR-01 to FR-02", "UC-01", "D-00", "Registration, authentication, role-aware routing"),
        ("FR-03 to FR-04", "UC-02", "D-01", "Profile CRUD and field-specific validation"),
        ("FR-05 to FR-17", "UC-03", "D-02, D-03", "Inputs, data retrieval, filters, calculations, price behavior"),
        ("FR-18 to FR-20", "UC-03, UC-04", "D-03, D-04", "Fixed strategies, explanation, comparison"),
        ("FR-21", "UC-05", "D-04", "Recalculation and immutable versioning"),
        ("FR-22 to FR-23", "UC-06", "D-03, D-04", "Manual refresh, change explanation, fallback"),
        ("FR-24", "UC-07", "D-05", "Report input, expiry, duplicate and image validation"),
        ("FR-25", "UC-08", "D-06", "Role restriction, decision reason, audit history"),
        ("NFR-01 to NFR-10", "All applicable", "All applicable", "Performance, failure, freshness, correctness, security, privacy, accessibility, usability, reliability"),
    ], widths=[1.25, 1.35, 1.25, 2.85], font_size=8)
    doc.add_heading("9.2 Acceptance Boundary", level=2)
    doc.add_paragraph("This version is the requirements and Use Case baseline supplied to the independent AI critique. The separately submitted AI Critique Report quotes the original critique response and identifies one selected finding. Changes made after this baseline must update the version number and repeat the critique if the submitted requirements or diagram no longer match.")
    doc.add_heading("9.3 Known Project Dependency", level=2)
    doc.add_paragraph("The live OneMap and LTA DataMall adapters must be smoke-tested before implementation claims are finalized. The MVP therefore requires deterministic Demo Fixture behavior and must label it explicitly. No requirement depends on unverified access to payment, reservation, charger control, or future-availability prediction.")

    props = doc.core_properties
    props.title = "PlugPlan SG Software Requirements Specification - Lab 1"
    props.subject = "SC2006 Lab 1 Requirements Elicitation"
    props.author = "Team SCE3-04"
    props.keywords = "SC2006, PlugPlan SG, SRS, requirements, use cases, UI mockups"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
